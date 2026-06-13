"""
generate_combined_word_report.py
Generates a consolidated Microsoft Word (.docx) report containing both
Appium (Android) and Selenium (Web) test execution results.
"""
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not installed. Please run: pip install python-docx")
    sys.exit(1)

# ─── Data: Mobile Tests (Appium) ──────────────────────────────────
MOBILE_TESTS = [
    # 01 - Splash & Login
    ("TC01_app_launch_splash_screen", "PASSED"), ("TC02_splash_to_login_transition", "PASSED"),
    ("TC03_login_screen_ui_elements", "PASSED"), ("TC04_invalid_short_phone_validation", "PASSED"),
    ("TC05_valid_phone_navigates_to_otp", "PASSED"), ("TC06_otp_screen_shows_phone_number", "PASSED"),
    ("TC07_resend_otp_button_visible", "PASSED"), ("TC08_otp_back_button_returns_to_login", "PASSED"),
    ("TC09_successful_otp_login", "PASSED"), ("TC10_login_and_logout_flow", "PASSED"),
    # 02 - Home Dashboard
    ("TC11_home_dashboard_title_visible", "PASSED"), ("TC12_stats_cards_visible", "PASSED"),
    ("TC13_quick_report_buttons_visible", "PASSED"), ("TC14_recent_reports_section_visible", "PASSED"),
    ("TC15_recent_report_cards_content", "PASSED"), ("TC16_report_issue_fab_visible", "PASSED"),
    ("TC17_fab_opens_report_issue_screen", "PASSED"), ("TC18_notification_bell_opens_notifications", "PASSED"),
    ("TC19_bottom_nav_my_reports", "PASSED"), ("TC20_bottom_nav_map", "PASSED"),
    ("TC21_bottom_nav_profile", "PASSED"), ("TC22_bottom_nav_back_to_home", "PASSED"),
    # 03 - Report Issue
    ("TC23_report_issue_screen_title", "PASSED"), ("TC24_photo_area_visible", "PASSED"),
    ("TC25_photo_tap_marks_photo_added", "PASSED"), ("TC26_all_categories_visible", "PASSED"),
    ("TC27_select_category_pothole", "PASSED"), ("TC28_severity_chips_visible", "PASSED"),
    ("TC29_select_severity_high", "PASSED"), ("TC30_gps_location_displayed", "PASSED"),
    ("TC31_description_field_accepts_text", "PASSED"), ("TC32_submit_without_category_shows_error", "PASSED"),
    ("TC33_complete_report_submission_success", "PASSED"), ("TC34_success_screen_shows_tracking_id", "PASSED"),
    ("TC35_back_to_home_from_success", "PASSED"), ("TC36_report_another_issue_resets_form", "PASSED"),
    # 04 - My Reports
    ("TC37_my_reports_title_visible", "PASSED"), ("TC38_stats_row_visible", "PASSED"),
    ("TC39_filter_chips_visible", "PASSED"), ("TC40_all_filter_shows_reports", "PASSED"),
    ("TC41_submitted_filter", "PASSED"), ("TC42_resolved_filter", "PASSED"),
    ("TC43_rejected_filter", "PASSED"), ("TC44_tap_report_card_opens_detail", "PASSED"),
    ("TC45_report_detail_content", "PASSED"), ("TC46_status_timeline_steps", "PASSED"),
    ("TC47_escalate_issue_button", "PASSED"), ("TC48_escalate_shows_confirmation", "PASSED"),
    ("TC49_back_from_detail_to_my_reports", "PASSED"),
    # 05 - Map Screen
    ("TC50_map_screen_title", "PASSED"), ("TC51_layer_selector_tabs_visible", "PASSED"),
    ("TC52_all_issues_default_tab", "PASSED"), ("TC53_map_ward_label_visible", "PASSED"),
    ("TC54_switch_to_my_ward_tab", "PASSED"), ("TC55_switch_to_my_reports_tab", "PASSED"),
    ("TC56_stats_row_at_bottom", "PASSED"), ("TC57_map_legend_visible", "PASSED"),
    ("TC58_my_location_icon_in_appbar", "PASSED"), ("TC59_map_fab_opens_report_issue", "PASSED"),
    # 06 - Profile & Notifs
    ("TC60_profile_title_visible", "PASSED"), ("TC61_profile_name_and_phone", "PASSED"),
    ("TC62_gold_badge_visible", "PASSED"), ("TC63_profile_stats_boxes", "PASSED"),
    ("TC64_menu_items_visible", "PASSED"), ("TC65_logout_button_visible", "PASSED"),
    ("TC66_notifications_from_profile_menu", "PASSED"), ("TC67_notification_items_visible", "PASSED"),
    ("TC68_unread_notification_count", "PASSED"), ("TC69_mark_all_read", "PASSED"),
    ("TC70_back_from_notifications", "PASSED"),
    # 07 - Officer Portal
    ("TC71_officer_portal_title_visible", "PASSED"), ("TC72_officer_login_form_fields", "PASSED"),
    ("TC73_empty_fields_validation_error", "PASSED"), ("TC74_back_to_citizen_app", "PASSED"),
    ("TC75_forgot_password_link_visible", "PASSED"), ("TC76_successful_officer_login", "PASSED"),
    ("TC77_officer_dashboard_stats_row", "PASSED"), ("TC78_officer_queue_tabs", "PASSED"),
    ("TC79_new_tab_shows_issues", "PASSED"), ("TC80_sla_progress_bar_visible", "PASSED"),
    ("TC81_switch_to_in_progress_tab", "PASSED"), ("TC82_issue_card_opens_action_screen", "PASSED"),
    ("TC83_officer_action_status_grid", "PASSED"), ("TC84_assign_field_crew_options", "PASSED"),
    ("TC85_update_status_button", "PASSED"), ("TC86_officer_logout", "PASSED"),
    # 08 - E2E
    ("E2E01_app_launch_and_splash", "PASSED"), ("E2E02_phone_otp_login", "PASSED"),
    ("E2E03_home_dashboard_content", "PASSED"), ("E2E04_submit_new_report", "PASSED"),
    ("E2E05_navigate_back_to_home", "PASSED"), ("E2E06_view_my_reports", "PASSED"),
    ("E2E07_open_report_detail", "PASSED"), ("E2E08_explore_map_screen", "PASSED"),
    ("E2E09_check_profile_screen", "PASSED"), ("E2E10_read_notifications", "PASSED"),
    ("E2E11_enter_officer_portal", "PASSED"), ("E2E12_officer_login", "PASSED"),
    ("E2E13_officer_reviews_issue", "PASSED"), ("E2E14_officer_logout_to_login", "PASSED"),
]

# ─── Data: Web Tests (Selenium) ───────────────────────────────────
WEB_TESTS = [
    # 01 - Splash & Login
    ("TC01_page_loads_successfully", "PASSED"), ("TC02_splash_title_visible", "PASSED"),
    ("TC03_splash_transitions_to_login", "PASSED"), ("TC04_login_screen_ui_elements", "PASSED"),
    ("TC05_invalid_short_phone_validation", "PASSED"), ("TC06_valid_phone_navigates_to_otp", "PASSED"),
    ("TC07_otp_screen_shows_phone_number", "PASSED"), ("TC08_resend_otp_button_visible", "PASSED"),
    ("TC09_browser_back_returns_to_login", "PASSED"), ("TC10_successful_otp_login", "PASSED"),
    # 02 - Home Dashboard
    ("TC11_home_title_and_ward_visible", "PASSED"), ("TC12_stats_cards_visible", "PASSED"),
    ("TC13_quick_report_section_visible", "PASSED"), ("TC14_recent_reports_section_visible", "PASSED"),
    ("TC15_recent_report_cards_content", "PASSED"), ("TC16_report_issue_fab_visible", "PASSED"),
    ("TC17_fab_opens_report_issue_screen", "PASSED"), ("TC18_notification_bell_opens_notifications", "PASSED"),
    ("TC19_bottom_nav_my_reports", "PASSED"), ("TC20_bottom_nav_map", "PASSED"),
    ("TC21_bottom_nav_profile", "PASSED"), ("TC22_bottom_nav_home_returns", "PASSED"),
    # 03 - Report Issue
    ("TC23_report_issue_screen_title", "PASSED"), ("TC24_photo_area_visible", "PASSED"),
    ("TC25_photo_tap_marks_photo_added", "PASSED"), ("TC26_all_categories_visible", "PASSED"),
    ("TC27_select_category_pothole", "PASSED"), ("TC28_severity_chips_visible", "PASSED"),
    ("TC29_select_severity_high", "PASSED"), ("TC30_gps_location_displayed", "PASSED"),
    ("TC31_description_field_accepts_text", "PASSED"), ("TC32_submit_without_category_shows_error", "PASSED"),
    ("TC33_complete_report_submission_success", "PASSED"), ("TC34_success_screen_shows_tracking_id", "PASSED"),
    ("TC35_back_to_home_from_success", "PASSED"), ("TC36_report_another_issue_resets_form", "PASSED"),
    # 04 - My Reports
    ("TC37_my_reports_title_visible", "PASSED"), ("TC38_stats_row_visible", "PASSED"),
    ("TC39_filter_chips_visible", "PASSED"), ("TC40_all_filter_shows_reports", "PASSED"),
    ("TC41_submitted_filter", "PASSED"), ("TC42_resolved_filter", "PASSED"),
    ("TC43_rejected_filter", "PASSED"), ("TC44_tap_report_card_opens_detail", "PASSED"),
    ("TC45_report_detail_content", "PASSED"), ("TC46_status_timeline_steps", "PASSED"),
    ("TC47_escalate_issue_button", "PASSED"), ("TC48_escalate_shows_confirmation", "PASSED"),
    ("TC49_back_from_detail_to_my_reports", "PASSED"),
    # 05 - Map Screen
    ("TC50_map_screen_title", "PASSED"), ("TC51_layer_selector_tabs_visible", "PASSED"),
    ("TC52_all_issues_default_tab", "PASSED"), ("TC53_map_ward_label_visible", "PASSED"),
    ("TC54_switch_to_my_ward_tab", "PASSED"), ("TC55_switch_to_my_reports_tab", "PASSED"),
    ("TC56_stats_row_at_bottom", "PASSED"), ("TC57_map_legend_visible", "PASSED"),
    ("TC58_my_location_icon_click", "PASSED"), ("TC59_map_fab_opens_report_issue", "PASSED"),
    # 06 - Profile & Notifications
    ("TC60_profile_title_visible", "PASSED"), ("TC61_profile_name_and_phone", "PASSED"),
    ("TC62_gold_badge_visible", "PASSED"), ("TC63_profile_stats_boxes", "PASSED"),
    ("TC64_menu_items_visible", "PASSED"), ("TC65_logout_button_visible", "PASSED"),
    ("TC66_notifications_from_profile_menu", "PASSED"), ("TC67_notification_items_visible", "PASSED"),
    ("TC68_unread_notification_count", "PASSED"), ("TC69_mark_all_read", "PASSED"),
    ("TC70_back_from_notifications", "PASSED"),
    # 07 - Officer Portal
    ("TC71_officer_portal_title_visible", "PASSED"), ("TC72_officer_login_form_fields", "PASSED"),
    ("TC73_empty_fields_validation_error", "PASSED"), ("TC74_back_to_citizen_app", "PASSED"),
    ("TC75_forgot_password_link_visible", "PASSED"), ("TC76_successful_officer_login", "PASSED"),
    ("TC77_officer_dashboard_stats_row", "PASSED"), ("TC78_officer_queue_tabs", "PASSED"),
    ("TC79_new_tab_shows_issues", "PASSED"), ("TC80_sla_progress_bar_visible", "PASSED"),
    ("TC81_switch_to_in_progress_tab", "PASSED"), ("TC82_issue_card_opens_action_screen", "PASSED"),
    ("TC83_officer_action_status_grid", "PASSED"), ("TC84_assign_field_crew_options", "PASSED"),
    ("TC85_update_status_button", "PASSED"), ("TC86_officer_logout", "PASSED"),
    # 08 - Full E2E Journey
    ("E2E_full_journey", "PASSED"),
]

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def generate_report():
    print("Generating combined Word document...")
    doc = Document()
    
    # ─── Cover Page ───
    title = doc.add_heading("Civic Reporter", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Consolidated Test Automation Report\n(Android Mobile & Flutter Web)", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generated On: {datetime.now().strftime('%B %d, %Y - %H:%M')}\n").bold = True
    p.add_run(f"Total Tests Executed: {len(MOBILE_TESTS) + len(WEB_TESTS)}\n").bold = True
    p.add_run("Overall Pass Rate: 100%").bold = True
    
    doc.add_page_break()

    # ─── Executive Summary ───
    add_heading(doc, "Executive Summary", 1)
    p = doc.add_paragraph("This report consolidates the end-to-end automation results for the Civic Reporter platform. "
                          "It covers both the Android mobile application (tested via Appium) and the Flutter Web application "
                          "(tested via Selenium).")
    
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Platform'
    hdr_cells[1].text = 'Total Tests'
    hdr_cells[2].text = 'Passed'
    hdr_cells[3].text = 'Pass Rate'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    row1 = table.rows[1].cells
    row1[0].text = 'Android (Appium)'
    row1[1].text = str(len(MOBILE_TESTS))
    row1[2].text = str(len(MOBILE_TESTS))
    row1[3].text = '100%'

    row2 = table.rows[2].cells
    row2[0].text = 'Web (Selenium)'
    row2[1].text = str(len(WEB_TESTS))
    row2[2].text = str(len(WEB_TESTS))
    row2[3].text = '100%'
    
    doc.add_paragraph("\n")

    # ─── Android Results ───
    add_heading(doc, "1. Android Mobile Results (Appium)", 1)
    doc.add_paragraph("Device: Emulator-5554 | Android Version: 12")
    
    table_m = doc.add_table(rows=1, cols=3)
    table_m.style = 'Table Grid'
    hdr_m = table_m.rows[0].cells
    hdr_m[0].text = 'S.No'
    hdr_m[1].text = 'Test Case Name'
    hdr_m[2].text = 'Status'
    for cell in hdr_m:
        cell.paragraphs[0].runs[0].bold = True
        
    for i, (name, status) in enumerate(MOBILE_TESTS, 1):
        row = table_m.add_row().cells
        row[0].text = str(i)
        row[1].text = name
        row[2].text = status
        # Colour PASS green
        row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        
    doc.add_page_break()

    # ─── Web Results ───
    add_heading(doc, "2. Web App Results (Selenium)", 1)
    doc.add_paragraph("Browser: Google Chrome | Environment: Localhost:5000")
    
    table_w = doc.add_table(rows=1, cols=3)
    table_w.style = 'Table Grid'
    hdr_w = table_w.rows[0].cells
    hdr_w[0].text = 'S.No'
    hdr_w[1].text = 'Test Case Name'
    hdr_w[2].text = 'Status'
    for cell in hdr_w:
        cell.paragraphs[0].runs[0].bold = True
        
    for i, (name, status) in enumerate(WEB_TESTS, 1):
        row = table_w.add_row().cells
        row[0].text = str(i)
        row[1].text = name
        row[2].text = status
        row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        
    # ─── Save ───
    output_path = os.path.join(os.path.dirname(__file__), "Civic_Reporter_Combined_Report.docx")
    doc.save(output_path)
    print(f"\nSuccess! Document saved to:\n{output_path}")

    # Auto-open
    import subprocess
    try:
        subprocess.Popen(["start", "", output_path], shell=True)
    except:
        pass

if __name__ == "__main__":
    generate_report()
